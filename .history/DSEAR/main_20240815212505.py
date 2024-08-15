import csv
from docx import Document
from docx.shared import Pt, RGBColor
import os
import re

# Function to read parameters from the CSV file
def read_parameters_from_csv(csv_file):
    params = []
    with open(csv_file, mode='r') as file:
        reader = csv.DictReader(file)
        for row in reader:
            # Parse the pages_to_keep and other parameters
            pages_to_keep = get_pages_to_keep(row["pages_to_keep"])
            row["pages_to_keep"] = pages_to_keep
            params.append(row)
    return params

# Function to map color names to RGBColor objects
def get_color(color_name):
    colors = {
        'red': RGBColor(255, 0, 0),
        'blue': RGBColor(0, 0, 255),
        'green': RGBColor(0, 128, 0),
        'orange': RGBColor(255, 165, 0),
        # Add more colors as needed
    }
    return colors.get(color_name.lower(), RGBColor(0, 0, 0))  # Default to black if color not found

# Function to safely convert font size
def get_font_size(font_size_value):
    try:
        return int(font_size_value) if font_size_value else 12
    except ValueError:
        return 12  # Default to 12 if conversion fails

# Function to get pages to keep from the CSV cell
def get_pages_to_keep(page_string):
    if not page_string:
        return []
    try:
        return list(map(int, page_string.strip('"').split(',')))
    except ValueError:
        return []

# Function to replace placeholders and apply formatting in a .docx file
def create_docx_with_replacements(template_docx_path, params, output_folder):
    for param in params:
        # Load the template document
        doc = Document(template_docx_path)
        
        # Extract only the pages specified for the current document
        pages_to_keep = param["pages_to_keep"]
        new_doc = Document()  # Create a new Document object to store the selected pages
        
        # Add only the paragraphs from the specified pages to new_doc
        for i, paragraph in enumerate(doc.paragraphs):
            if i in pages_to_keep:  # Assuming paragraph index corresponds to page number (This may need adjustment)
                new_doc.add_paragraph(paragraph.text, style=paragraph.style)
                
                # Apply formatting to paragraphs
                for key in param:
                    match = re.match(r'xxx(\d+)_text', key)
                    if match:
                        placeholder_number = match.group(1)
                        placeholder = f"xxx{placeholder_number}"
                        text_key = key
                        font_size_key = f"xxx{placeholder_number}_font_size"
                        color_key = f"xxx{placeholder_number}_color"

                        # Skip this placeholder if any of the parameters are None
                        if not param.get(text_key) or param[text_key].lower() == 'none':
                            continue

                        # Only attempt replacement if the placeholder is found in the text
                        if placeholder in paragraph.text:
                            replace_text_with_formatting(
                                new_doc.paragraphs[-1],
                                placeholder,
                                param[text_key],
                                get_font_size(param.get(font_size_key)),  # Safely get font size
                                get_color(param.get(color_key, 'black'))  # Default color is black if not specified
                            )

        # Save the new document in the specified output folder
        output_filename = os.path.join(output_folder, param["output_file"])
        new_doc.save(output_filename)
        print(f"{output_filename} created successfully.")

# Function to replace text and apply formatting
def replace_text_with_formatting(paragraph, placeholder, new_text, font_size, font_color):
    # Replace the placeholder while preserving formatting
    if placeholder in paragraph.text:
        inline = paragraph.runs
        for run in inline:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, new_text)
                run.font.size = Pt(font_size)
                run.font.color.rgb = font_color

# Read parameters from CSV
params = read_parameters_from_csv('parameters.csv')

# Define the template path and output folder
template_path = os.path.join('docs', 'template.docx')
output_folder = 'docs/generated'

# Use the function with your template DOCX and parameters
create_docx_with_replacements(template_path, params, output_folder)
