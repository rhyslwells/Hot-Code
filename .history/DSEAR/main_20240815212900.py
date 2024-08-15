import csv
from docx import Document
from docx.shared import Pt, RGBColor
import os
import re

# Function to read parameters from the CSV file
def read_parameters_from_csv(csv_file):
    params = []
    with open(csv_file, mode='r') as file:
        reader = csv.DictReader(file)
        for row in reader:
            # Parse the pages_to_keep and other parameters
            pages_to_keep = get_pages_to_keep(row["pages_to_keep"])
            row["pages_to_keep"] = pages_to_keep
            params.append(row)
    return params

# Function to map color names to RGBColor objects
def get_color(color_name):
    colors = {
        'red': RGBColor(255, 0, 0),
        'blue': RGBColor(0, 0, 255),
        'green': RGBColor(0, 128, 0),
        'orange': RGBColor(255, 165, 0),
        # Add more colors as needed
    }
    return colors.get(color_name.lower(), RGBColor(0, 0, 0))  # Default to black if color not found

# Function to safely convert font size
def get_font_size(font_size_value):
    try:
        return int(font_size_value) if font_size_value else 12
    except ValueError:
        return 12  # Default to 12 if conversion fails

# Function to get pages to keep from the CSV cell
def get_pages_to_keep(page_string):
    if not page_string:
        return []
    try:
        return list(map(int, page_string.strip('"').split(',')))
    except ValueError:
        return []

# Function to create a new document with only the specified pages
def create_filtered_docx(original_doc_path, pages_to_keep, output_doc_path):
    doc = Document(original_doc_path)
    new_doc = Document()

    page_indices = find_page_indices(doc)

    for page_number in pages_to_keep:
        if page_number <= len(page_indices):
            start_index = page_indices[page_number - 1][0]
            end_index = page_indices[page_number - 1][1]

            for i in range(start_index, end_index):
                new_doc.add_paragraph(doc.paragraphs[i].text, style=doc.paragraphs[i].style)

    new_doc.save(output_doc_path)

def find_page_indices(doc):
    """ This function should return a list of tuples indicating the start and end paragraph indices of each page. """
    indices = []
    start_index = 0
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == '':
            # Assuming page breaks or empty paragraphs denote page ends
            indices.append((start_index, i))
            start_index = i + 1
    # Add the last section
    indices.append((start_index, len(doc.paragraphs)))
    return indices

# Function to replace placeholders and apply formatting in a .docx file
def create_docx_with_replacements(template_docx_path, params, output_folder):
    for param in params:
        output_filename = os.path.join(output_folder, param["output_file"])
        pages_to_keep = param["pages_to_keep"]

        # Create filtered document with only the specified pages
        create_filtered_docx(template_docx_path, pages_to_keep, output_filename)

        # Load the filtered document for replacement and formatting
        doc = Document(output_filename)

        # Apply formatting to paragraphs
        for key in param:
            match = re.match(r'xxx(\d+)_text', key)
            if match:
                placeholder_number = match.group(1)
                placeholder = f"xxx{placeholder_number}"
                text_key = key
                font_size_key = f"xxx{placeholder_number}_font_size"
                color_key = f"xxx{placeholder_number}_color"

                # Skip this placeholder if any of the parameters are None
                if not param.get(text_key) or param[text_key].lower() == 'none':
                    continue

                for paragraph in doc.paragraphs:
                    if placeholder in paragraph.text:
                        replace_text_with_formatting(
                            paragraph,
                            placeholder,
                            param[text_key],
                            get_font_size(param.get(font_size_key)),  # Safely get font size
                            get_color(param.get(color_key, 'black'))  # Default color is black if not specified
                        )

        # Save the formatted document
        doc.save(output_filename)
        print(f"{output_filename} created successfully.")

# Function to replace text and apply formatting
def replace_text_with_formatting(paragraph, placeholder, new_text, font_size, font_color):
    # Replace the placeholder while preserving formatting
    if placeholder in paragraph.text:
        inline = paragraph.runs
        for run in inline:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, new_text)
                run.font.size = Pt(font_size)
                run.font.color.rgb = font_color

# Read parameters from CSV
params = read_parameters_from_csv('parameters.csv')

# Define the template path and output folder
template_path = os.path.join('docs', 'template.docx')
output_folder = 'docs/generated'

# Use the function with your template DOCX and parameters
create_docx_with_replacements(template_path, params, output_folder)
